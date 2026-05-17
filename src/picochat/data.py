"""Dataset inspection and simple corpus building utilities."""

from __future__ import annotations

from dataclasses import asdict, dataclass
from fnmatch import fnmatch
import json
from pathlib import Path
import re
import zlib

from picochat.dataset_pack import DatasetPack, load_dataset_pack
from picochat.tuning_data import (
    ChatEvalDataReport,
    ChatSFTDataReport,
    inspect_chat_eval_data,
    inspect_chat_sft_data,
)


TEXT_EXTENSIONS = {
    ".md",
    ".txt",
    ".text",
    ".jsonl",
    ".csv",
    ".py",
}
CORPUS_REPORT_MAX_DOCUMENT_ROWS = 200
CORPUS_REPORT_MAX_FILE_ROWS = 500

DOCUMENT_EXTENSIONS = {
    ".docx",
    ".pdf",
}

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS
DEFAULT_CHAT_INPUT = "examples/tiny_chat.jsonl"
DEFAULT_EVAL_INPUT = "examples/tiny_eval.jsonl"
DEFAULT_CORPUS_INPUT = "examples/tiny_corpus.txt"
NEAR_DUPLICATE_THRESHOLD = 0.82
NEAR_DUPLICATE_MAX_DOCUMENTS = 50_000
NEAR_DUPLICATE_MAX_TOKENS = 768
NEAR_DUPLICATE_SIGNATURE_SIZE = 32
NEAR_DUPLICATE_BAND_SIZE = 4
NEAR_DUPLICATE_BUCKET_LIMIT = 64
_NEAR_DUPLICATE_TOKEN_RE = re.compile(r"[a-z0-9]+")


class DocumentExtractionError(RuntimeError):
    """Raised when a source document cannot be converted into training text."""

    def __init__(self, reason: str) -> None:
        super().__init__(reason)
        self.reason = reason


@dataclass(frozen=True)
class CorpusStats:
    num_files: int
    num_documents: int
    num_characters: int
    num_lines: int
    average_document_chars: float
    duplicate_document_rate: float
    duplicate_line_rate: float
    non_ascii_rate: float
    empty_line_rate: float
    near_duplicate_document_rate: float = 0.0
    near_duplicate_document_pairs: int = 0
    near_duplicate_documents_checked: int = 0

    def to_dict(self) -> dict[str, float | int]:
        return asdict(self)


@dataclass(frozen=True)
class CorpusFileRecord:
    path: str
    extension: str
    num_characters: int
    num_lines: int
    included: bool
    reason: str
    quality_score: int
    quality_flags: tuple[str, ...]
    label: str | None = None

    def to_dict(self) -> dict[str, object]:
        return asdict(self)


@dataclass(frozen=True)
class CorpusDocumentRecord:
    document_id: int
    path: str
    label: str | None
    char_start: int
    char_end: int
    num_characters: int
    num_lines: int
    quality_score: int
    quality_flags: tuple[str, ...]

    def to_dict(self) -> dict[str, object]:
        return asdict(self)


@dataclass(frozen=True)
class CorpusReadinessCheck:
    name: str
    status: str
    metric: str
    threshold: str
    message: str

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass(frozen=True)
class CorpusReadiness:
    status: str
    summary: str
    checks: tuple[CorpusReadinessCheck, ...]

    def to_dict(self) -> dict:
        return {
            "status": self.status,
            "summary": self.summary,
            "checks": [check.to_dict() for check in self.checks],
        }


@dataclass(frozen=True)
class CorpusTrainingBudget:
    preset: str
    estimated_tokens: int
    suggested_context_size: int
    estimated_windows: int
    suggested_batch_size: int
    suggested_base_steps: int
    estimated_tokens_per_step: int
    estimated_passes: float
    note: str

    def to_dict(self) -> dict[str, float | int | str]:
        return asdict(self)


@dataclass(frozen=True)
class CorpusTrainingCommand:
    out_dir: str
    chat_input: str
    eval_input: str
    dataset_pack: str | None
    command: str
    note: str

    def to_dict(self) -> dict[str, str | None]:
        return asdict(self)


@dataclass(frozen=True)
class CorpusBuildReport:
    input_path: str
    output_path: str
    manifest_path: str
    report_path: str
    stats: CorpusStats
    files: tuple[CorpusFileRecord, ...]
    readiness: CorpusReadiness
    budget: CorpusTrainingBudget
    training_command: CorpusTrainingCommand
    chat_data: ChatSFTDataReport
    eval_data: ChatEvalDataReport
    warnings: tuple[str, ...]
    recipe_path: str | None = None
    dataset_pack: str | None = None
    min_quality_score: int = 0
    documents: tuple[CorpusDocumentRecord, ...] = ()

    def to_dict(self) -> dict:
        return {
            "input_path": self.input_path,
            "output_path": self.output_path,
            "manifest_path": self.manifest_path,
            "report_path": self.report_path,
            "recipe_path": self.recipe_path,
            "dataset_pack": self.dataset_pack,
            "min_quality_score": self.min_quality_score,
            "documents": [document.to_dict() for document in self.documents],
            "stats": self.stats.to_dict(),
            "files": [record.to_dict() for record in self.files],
            "readiness": self.readiness.to_dict(),
            "budget": self.budget.to_dict(),
            "training_command": self.training_command.to_dict(),
            "chat_data": self.chat_data.to_dict(),
            "eval_data": self.eval_data.to_dict(),
            "warnings": list(self.warnings),
        }


@dataclass(frozen=True)
class CorpusPreviewReport:
    input_path: str
    recipe_path: str | None
    dataset_pack: str | None
    stats: CorpusStats
    files: tuple[CorpusFileRecord, ...]
    readiness: CorpusReadiness
    budget: CorpusTrainingBudget
    training_command: CorpusTrainingCommand
    chat_data: ChatSFTDataReport
    eval_data: ChatEvalDataReport
    warnings: tuple[str, ...]
    preview: str
    min_quality_score: int = 0

    def to_dict(self) -> dict:
        return {
            "input_path": self.input_path,
            "recipe_path": self.recipe_path,
            "dataset_pack": self.dataset_pack,
            "min_quality_score": self.min_quality_score,
            "stats": self.stats.to_dict(),
            "files": [record.to_dict() for record in self.files],
            "readiness": self.readiness.to_dict(),
            "budget": self.budget.to_dict(),
            "training_command": self.training_command.to_dict(),
            "chat_data": self.chat_data.to_dict(),
            "eval_data": self.eval_data.to_dict(),
            "warnings": list(self.warnings),
            "preview": self.preview,
        }


@dataclass(frozen=True)
class _CollectedCorpus:
    input_path: str
    recipe_path: str | None
    dataset_pack: str | None
    documents: tuple[str, ...]
    document_sources: tuple[CorpusFileRecord, ...]
    files: tuple[CorpusFileRecord, ...]
    stats: CorpusStats
    readiness: CorpusReadiness
    budget: CorpusTrainingBudget
    training_command: CorpusTrainingCommand
    chat_data: ChatSFTDataReport
    eval_data: ChatEvalDataReport
    warnings: tuple[str, ...]
    min_quality_score: int


@dataclass(frozen=True)
class _SourceCandidate:
    path: Path
    label: str | None = None
    skipped_reason: str | None = None


def find_text_files(path: str | Path) -> list[Path]:
    """Return readable text-like files under a file or directory path."""
    root = Path(path)
    if root.is_file():
        return [root]
    if not root.exists():
        raise FileNotFoundError(root)
    files = [
        item for item in root.rglob("*")
        if item.is_file() and item.suffix.lower() in TEXT_EXTENSIONS
    ]
    return sorted(files)


def find_corpus_files(path: str | Path) -> list[Path]:
    """Return supported corpus source files under a file or directory path."""
    root = Path(path)
    if root.is_file():
        return [root]
    if not root.exists():
        raise FileNotFoundError(root)
    files = [
        item for item in root.rglob("*")
        if item.is_file() and item.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    return sorted(files)


def read_documents(path: str | Path) -> list[str]:
    """Read one document per supported corpus source file."""
    documents: list[str] = []
    for file_path in find_corpus_files(path):
        try:
            text, _reason = _extract_source_text(file_path)
        except DocumentExtractionError:
            continue
        if text.strip():
            documents.append(text)
    return documents


def read_corpus_documents(
    input_path: str | Path | None = None,
    recipe_path: str | Path | None = None,
    *,
    dataset_pack: str | Path | None = None,
    min_quality_score: int = 0,
) -> tuple[str, list[str]]:
    """Read corpus documents from a direct path, recipe, or dataset pack."""
    if dataset_pack or recipe_path:
        collected = _collect_corpus_sources(
            input_path,
            recipe_path,
            dataset_pack=dataset_pack,
            min_quality_score=min_quality_score,
        )
        return collected.input_path, list(collected.documents)
    if input_path is None:
        raise ValueError("input_path, recipe_path, or dataset_pack is required")
    return str(input_path), read_documents(input_path)


def inspect_documents(documents: list[str], num_files: int | None = None) -> CorpusStats:
    """Compute simple, explainable corpus quality stats."""
    if not documents:
        return CorpusStats(
            num_files=num_files or 0,
            num_documents=0,
            num_characters=0,
            num_lines=0,
            average_document_chars=0.0,
            duplicate_document_rate=0.0,
            duplicate_line_rate=0.0,
            non_ascii_rate=0.0,
            empty_line_rate=0.0,
            near_duplicate_document_rate=0.0,
            near_duplicate_document_pairs=0,
            near_duplicate_documents_checked=0,
        )

    seen_lines: set[str] = set()
    duplicate_lines = 0
    non_empty_line_count = 0
    empty_lines = 0
    num_lines = 0
    non_ascii_chars = 0
    document_chars = 0
    for document in documents:
        document_chars += len(document)
        non_ascii_chars += sum(1 for char in document if ord(char) > 127)
        lines = document.splitlines()
        num_lines += len(lines)
        for line in lines:
            stripped = line.strip()
            if not stripped:
                empty_lines += 1
                continue
            non_empty_line_count += 1
            if stripped in seen_lines:
                duplicate_lines += 1
            else:
                seen_lines.add(stripped)
    joined_character_count = document_chars + max(0, len(documents) - 1)
    normalized_documents = [_normalize_document_for_duplicate_check(document) for document in documents]
    duplicate_documents = len(normalized_documents) - len(set(normalized_documents))
    near_duplicate_rate, near_duplicate_pairs, near_duplicate_checked = (
        _near_duplicate_document_stats(documents)
    )

    return CorpusStats(
        num_files=num_files if num_files is not None else len(documents),
        num_documents=len(documents),
        num_characters=joined_character_count,
        num_lines=num_lines,
        average_document_chars=joined_character_count / len(documents),
        duplicate_document_rate=duplicate_documents / max(1, len(normalized_documents)),
        duplicate_line_rate=duplicate_lines / max(1, non_empty_line_count),
        non_ascii_rate=non_ascii_chars / max(1, joined_character_count),
        empty_line_rate=empty_lines / max(1, num_lines),
        near_duplicate_document_rate=near_duplicate_rate,
        near_duplicate_document_pairs=near_duplicate_pairs,
        near_duplicate_documents_checked=near_duplicate_checked,
    )


def inspect_path(path: str | Path) -> CorpusStats:
    files = find_corpus_files(path)
    documents = read_documents(path)
    return inspect_documents(documents, num_files=len(files))


def _normalize_document_for_duplicate_check(text: str) -> str:
    return " ".join(text.lower().split())


def _near_duplicate_document_stats(documents: list[str]) -> tuple[float, int, int]:
    """Estimate near-duplicate document rate with deterministic bottom-k MinHash LSH."""
    if len(documents) < 2:
        return 0.0, 0, len(documents)

    selected_indices = _evenly_spaced_indices(len(documents), NEAR_DUPLICATE_MAX_DOCUMENTS)
    signatures: list[frozenset[int]] = []
    buckets: dict[tuple[int, tuple[int, ...]], list[int]] = {}
    near_duplicate_docs: set[int] = set()
    near_duplicate_pairs = 0
    seen_pairs: set[tuple[int, int]] = set()

    for selected_position, doc_index in enumerate(selected_indices):
        signature = _document_minhash_signature(documents[doc_index])
        signatures.append(signature)
        if len(signature) < NEAR_DUPLICATE_BAND_SIZE * 2:
            continue
        ordered = tuple(sorted(signature))
        for band_start in range(0, len(ordered), NEAR_DUPLICATE_BAND_SIZE):
            band = ordered[band_start:band_start + NEAR_DUPLICATE_BAND_SIZE]
            if len(band) < NEAR_DUPLICATE_BAND_SIZE:
                continue
            bucket_key = (band_start, band)
            bucket = buckets.setdefault(bucket_key, [])
            for other_position in bucket[-NEAR_DUPLICATE_BUCKET_LIMIT:]:
                pair = (other_position, selected_position)
                if pair in seen_pairs:
                    continue
                seen_pairs.add(pair)
                if _signature_jaccard(signatures[other_position], signature) >= NEAR_DUPLICATE_THRESHOLD:
                    near_duplicate_pairs += 1
                    near_duplicate_docs.add(other_position)
                    near_duplicate_docs.add(selected_position)
            bucket.append(selected_position)

    checked = len(selected_indices)
    return len(near_duplicate_docs) / max(1, checked), near_duplicate_pairs, checked


def _evenly_spaced_indices(total: int, limit: int) -> list[int]:
    if total <= limit:
        return list(range(total))
    return sorted({round(index * (total - 1) / (limit - 1)) for index in range(limit)})


def _document_minhash_signature(text: str) -> frozenset[int]:
    tokens = _NEAR_DUPLICATE_TOKEN_RE.findall(text.lower())[:NEAR_DUPLICATE_MAX_TOKENS]
    if len(tokens) < 5:
        return frozenset()
    hashes = {
        zlib.crc32(" ".join(tokens[index:index + 5]).encode("utf-8"))
        for index in range(len(tokens) - 4)
    }
    if len(hashes) <= NEAR_DUPLICATE_SIGNATURE_SIZE:
        return frozenset(hashes)
    return frozenset(sorted(hashes)[:NEAR_DUPLICATE_SIGNATURE_SIZE])


def _signature_jaccard(left: frozenset[int], right: frozenset[int]) -> float:
    if not left or not right:
        return 0.0
    return len(left & right) / len(left | right)


def assess_corpus_readiness(
    stats: CorpusStats,
    records: list[CorpusFileRecord] | tuple[CorpusFileRecord, ...],
) -> CorpusReadiness:
    """Return structured, explainable readiness checks for tiny training."""
    skipped = [record for record in records if not record.included]
    checks = [
        _readiness_check(
            "usable_documents",
            "fail" if stats.num_documents == 0 else "pass",
            str(stats.num_documents),
            ">= 1",
            "At least one supported source needs usable text.",
        ),
        _readiness_check(
            "corpus_size",
            "fail" if stats.num_characters == 0 else "warn" if stats.num_characters < 1000 else "pass",
            f"{stats.num_characters:,} chars",
            ">= 1,000 chars",
            "Very small corpora are useful for smoke tests, but they mostly teach memorization.",
        ),
        _readiness_check(
            "document_mix",
            "warn" if stats.num_documents == 1 else "pass" if stats.num_documents > 1 else "fail",
            str(stats.num_documents),
            ">= 2 for broader training",
            "One document is fine for a focused overfit test; more documents give better variation.",
        ),
        _readiness_check(
            "duplicate_documents",
            "warn" if stats.duplicate_document_rate > 0.05 else "pass",
            f"{stats.duplicate_document_rate * 100:.2f}%",
            "<= 5%",
            "Repeated full documents make eval and samples look better than the model really is.",
        ),
        _readiness_check(
            "duplicate_lines",
            "warn" if stats.duplicate_line_rate > 0.15 else "pass",
            f"{stats.duplicate_line_rate * 100:.2f}%",
            "<= 15%",
            "Repeated lines can make a tiny model memorize phrasing instead of learning patterns.",
        ),
        _readiness_check(
            "near_duplicate_documents",
            "warn" if stats.near_duplicate_document_rate > 0.05 else "pass",
            f"{stats.near_duplicate_document_rate * 100:.2f}% over "
            f"{stats.near_duplicate_documents_checked:,} checked",
            "<= 5%",
            "Near-duplicate sources reduce effective data diversity even when exact dedup looks clean.",
        ),
        _readiness_check(
            "empty_lines",
            "warn" if stats.empty_line_rate > 0.35 else "pass",
            f"{stats.empty_line_rate * 100:.2f}%",
            "<= 35%",
            "Too many empty lines waste context windows during next-token training.",
        ),
        _readiness_check(
            "non_ascii",
            "warn" if stats.non_ascii_rate > 0.05 else "pass",
            f"{stats.non_ascii_rate * 100:.2f}%",
            "<= 5% unless intentional",
            "High non-ASCII text is fine when expected; otherwise it may indicate extraction noise.",
        ),
        _readiness_check(
            "skipped_sources",
            "warn" if skipped else "pass",
            str(len(skipped)),
            "0 skipped preferred",
            "Skipped files may mean unsupported formats, missing extractors, or recipe exclusions.",
        ),
    ]
    if any(check.status == "fail" for check in checks):
        status = "blocked"
        summary = "Corpus is not trainable yet."
    elif any(check.status == "warn" for check in checks):
        status = "caution"
        summary = "Corpus is trainable, but read the cautions before spending time on training."
    else:
        status = "ready"
        summary = "Corpus looks ready for a tiny training run."
    return CorpusReadiness(status=status, summary=summary, checks=tuple(checks))


def estimate_training_budget(stats: CorpusStats) -> CorpusTrainingBudget:
    """Estimate a conservative first training budget before tokenizer training."""
    estimated_tokens = stats.num_characters
    if estimated_tokens == 0:
        context_size = 32
        batch_size = 1
        base_steps = 0
        preset = "blocked"
        note = "No usable text means there is nothing to train yet."
    elif estimated_tokens < 1000:
        context_size = 32
        batch_size = 4
        base_steps = 100
        preset = "smoke"
        note = "Use this only to test the pipeline; the model will mostly memorize."
    elif stats.num_documents == 1:
        context_size = 128
        batch_size = 8
        base_steps = 300
        preset = "overfit-check"
        note = "Good for checking whether the model can learn one source before adding more documents."
    elif estimated_tokens < 50000:
        context_size = 128
        batch_size = 8
        base_steps = 500
        preset = "tiny"
        note = "Reasonable first tiny run; compare loss, samples, and eval before increasing scale."
    else:
        context_size = 256
        batch_size = 4
        base_steps = 1000
        preset = "small-preview"
        note = "Large enough for a longer local run; start here before trying larger models."

    estimated_windows = max(0, estimated_tokens - context_size)
    tokens_per_step = context_size * batch_size
    estimated_passes = (base_steps * tokens_per_step / estimated_tokens) if estimated_tokens else 0.0
    return CorpusTrainingBudget(
        preset=preset,
        estimated_tokens=estimated_tokens,
        suggested_context_size=context_size,
        estimated_windows=estimated_windows,
        suggested_batch_size=batch_size,
        suggested_base_steps=base_steps,
        estimated_tokens_per_step=tokens_per_step,
        estimated_passes=estimated_passes,
        note=note,
    )


def suggest_training_command(
    input_path: str,
    recipe_path: str | None,
    budget: CorpusTrainingBudget,
    chat_input: str | None = None,
    eval_input: str | None = None,
    dataset_pack: str | None = None,
    min_quality_score: int = 0,
) -> CorpusTrainingCommand:
    """Build a copyable first-run command from corpus intake metadata."""
    out_dir = f"runs/{_slugify_path(dataset_pack or recipe_path or input_path)}-v1"
    chat_input = _default_path(chat_input, DEFAULT_CHAT_INPUT)
    eval_input = _default_path(eval_input, DEFAULT_EVAL_INPUT)
    if budget.preset == "blocked":
        return CorpusTrainingCommand(
            out_dir=out_dir,
            chat_input=chat_input,
            eval_input=eval_input,
            dataset_pack=dataset_pack,
            command="",
            note="No training command yet; fix blocked corpus readiness checks first.",
        )

    custom_source = not dataset_pack and not (
        recipe_path is None and _same_path_text(input_path, DEFAULT_CORPUS_INPUT)
    )
    uses_default_tuning = chat_input == DEFAULT_CHAT_INPUT or eval_input == DEFAULT_EVAL_INPUT
    if custom_source and uses_default_tuning:
        return CorpusTrainingCommand(
            out_dir=out_dir,
            chat_input=chat_input,
            eval_input=eval_input,
            dataset_pack=dataset_pack,
            command="",
            note=(
                "No training command yet; create domain chat/eval files or a dataset pack. "
                "Picochat will not silently train a custom corpus with demo tuning data."
            ),
        )

    if dataset_pack:
        source_args = ["--dataset-pack", dataset_pack]
    else:
        source_flag = "--corpus-recipe" if recipe_path else "--corpus-input"
        source_args = [
            source_flag,
            recipe_path or input_path,
            "--chat-input",
            chat_input,
            "--eval-input",
            eval_input,
        ]
    command = _shell_command([
        "PYTHONPATH=src",
        "python",
        "-m",
        "picochat.cli",
        "run",
        "tiny",
        "--out-dir",
        out_dir,
        *source_args,
        *(["--scale", "pico"] if budget.preset == "small-preview" else []),
        "--context-size",
        budget.suggested_context_size,
        "--base-batch-size",
        budget.suggested_batch_size,
        "--base-steps",
        budget.suggested_base_steps,
        "--base-early-stop-patience",
        3,
        "--sft-early-stop-patience",
        4,
        "--canary-count",
        3,
        "--sft-sampling",
        "category_sqrt",
        "--split-mode",
        "document",
        *(["--min-score", min_quality_score] if min_quality_score else []),
    ])
    note = "Uses the selected chat/eval JSONL files for SFT and scoring."
    if dataset_pack:
        note = "Uses the dataset pack corpus, chat SFT, and eval files."
    elif chat_input == DEFAULT_CHAT_INPUT and eval_input == DEFAULT_EVAL_INPUT:
        note = "Uses default chat/eval examples; replace --chat-input and --eval-input for domain-specific tuning."
    elif chat_input == DEFAULT_CHAT_INPUT or eval_input == DEFAULT_EVAL_INPUT:
        note = "One tuning file is still using a default example; replace it before a real domain run."
    return CorpusTrainingCommand(
        out_dir=out_dir,
        chat_input=chat_input,
        eval_input=eval_input,
        dataset_pack=dataset_pack,
        command=command,
        note=note,
    )


def _slugify_path(path: str) -> str:
    stem = Path(path).stem or "corpus"
    slug = "".join(char.lower() if char.isalnum() else "-" for char in stem).strip("-")
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug or "corpus"


def _default_path(path: str | Path | None, default: str) -> str:
    text = str(path).strip() if path is not None else ""
    return text or default


def _same_path_text(left: str | Path | None, right: str) -> bool:
    return str(left or "").strip() == right


def _shell_command(parts: list[object]) -> str:
    return " ".join(_shell_token(part) for part in parts)


def _shell_token(value: object) -> str:
    text = str(value)
    if text and all(char.isalnum() or char in "_./:=+-" for char in text):
        return text
    return "'" + text.replace("'", "'\\''") + "'"


def _readiness_check(name: str, status: str, metric: str, threshold: str, message: str) -> CorpusReadinessCheck:
    return CorpusReadinessCheck(
        name=name,
        status=status,
        metric=metric,
        threshold=threshold,
        message=message,
    )


def build_corpus(input_path: str | Path, output_path: str | Path) -> CorpusStats:
    """Combine corpus sources into one normalized text file."""
    return build_corpus_artifacts(input_path, output_path, write_manifest=False).stats


def build_corpus_artifacts(
    input_path: str | Path | None,
    output_path: str | Path,
    manifest_path: str | Path | None = None,
    report_path: str | Path | None = None,
    write_manifest: bool = True,
    recipe_path: str | Path | None = None,
    chat_input: str | Path | None = None,
    eval_input: str | Path | None = None,
    dataset_pack: str | Path | None = None,
    min_quality_score: int = 0,
) -> CorpusBuildReport:
    """Combine corpus sources and write provenance artifacts."""
    output_path = Path(output_path)
    manifest_path = Path(manifest_path) if manifest_path else output_path.with_name("corpus_manifest.json")
    report_path = Path(report_path) if report_path else output_path.with_name("corpus_report.md")
    collected = _collect_corpus_sources(
        input_path,
        recipe_path,
        chat_input,
        eval_input,
        dataset_pack,
        min_quality_score=min_quality_score,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    _write_corpus_documents(output_path, collected.documents)
    document_records = _document_records(collected.document_sources, collected.documents)

    report = CorpusBuildReport(
        input_path=collected.input_path,
        output_path=str(output_path),
        manifest_path=str(manifest_path),
        report_path=str(report_path),
        stats=collected.stats,
        files=collected.files,
        readiness=collected.readiness,
        budget=collected.budget,
        training_command=collected.training_command,
        chat_data=collected.chat_data,
        eval_data=collected.eval_data,
        warnings=collected.warnings,
        recipe_path=collected.recipe_path,
        dataset_pack=collected.dataset_pack,
        min_quality_score=collected.min_quality_score,
        documents=document_records,
    )
    if write_manifest:
        manifest_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.parent.mkdir(parents=True, exist_ok=True)
        manifest_path.write_text(json.dumps(report.to_dict(), indent=2), encoding="utf-8")
        report_path.write_text(corpus_report_markdown(report), encoding="utf-8")
    return report


def preview_corpus_sources(
    input_path: str | Path | None = None,
    recipe_path: str | Path | None = None,
    preview_chars: int = 1000,
    chat_input: str | Path | None = None,
    eval_input: str | Path | None = None,
    dataset_pack: str | Path | None = None,
    min_quality_score: int = 0,
) -> CorpusPreviewReport:
    """Inspect corpus sources without writing artifacts."""
    collected = _collect_corpus_sources(
        input_path,
        recipe_path,
        chat_input,
        eval_input,
        dataset_pack,
        min_quality_score=min_quality_score,
    )
    return CorpusPreviewReport(
        input_path=collected.input_path,
        recipe_path=collected.recipe_path,
        dataset_pack=collected.dataset_pack,
        stats=collected.stats,
        files=collected.files,
        readiness=collected.readiness,
        budget=collected.budget,
        training_command=collected.training_command,
        chat_data=collected.chat_data,
        eval_data=collected.eval_data,
        warnings=collected.warnings,
        preview=_preview_documents(collected.documents, max(0, preview_chars)),
        min_quality_score=collected.min_quality_score,
    )


def corpus_report_markdown(report: CorpusBuildReport) -> str:
    """Render a human-readable corpus provenance report."""
    included = [record for record in report.files if record.included]
    skipped = [record for record in report.files if not record.included]
    stats = report.stats
    lines = [
        "# Picochat Corpus Report",
        "",
        "This report records which local files were used to build the training corpus.",
        "",
        "## Summary",
        "",
        f"- Input path: `{report.input_path}`",
        f"- Output corpus: `{report.output_path}`",
        f"- Recipe: `{report.recipe_path}`" if report.recipe_path else "- Recipe: none",
        f"- Dataset pack: `{report.dataset_pack}`" if report.dataset_pack else "- Dataset pack: none",
        f"- Minimum source quality score: {report.min_quality_score}",
        f"- Files scanned: {len(report.files)}",
        f"- Files included: {len(included)}",
        f"- Files skipped: {len(skipped)}",
        f"- Documents: {stats.num_documents}",
        f"- Characters: {stats.num_characters:,}",
        f"- Lines: {stats.num_lines:,}",
        f"- Duplicate document rate: {stats.duplicate_document_rate * 100:.2f}%",
        f"- Near-duplicate document rate: {stats.near_duplicate_document_rate * 100:.2f}% "
        f"({stats.near_duplicate_document_pairs:,} pair(s), "
        f"{stats.near_duplicate_documents_checked:,} checked)",
        f"- Duplicate line rate: {stats.duplicate_line_rate * 100:.2f}%",
        f"- Empty line rate: {stats.empty_line_rate * 100:.2f}%",
        f"- Non-ASCII rate: {stats.non_ascii_rate * 100:.2f}%",
        "",
        "## Readiness",
        "",
        f"- Status: `{report.readiness.status}`",
        f"- Summary: {report.readiness.summary}",
        "",
        "| Check | Status | Metric | Threshold | Note |",
        "| --- | --- | ---: | --- | --- |",
    ]
    for check in report.readiness.checks:
        lines.append(
            f"| `{check.name}` | `{check.status}` | {check.metric} | `{check.threshold}` | {check.message} |"
        )
    budget = report.budget
    lines.extend([
        "",
        "## Training Budget",
        "",
        f"- Preset: `{budget.preset}`",
        f"- Estimated char tokens: {budget.estimated_tokens:,}",
        f"- Suggested context size: {budget.suggested_context_size}",
        f"- Estimated training windows: {budget.estimated_windows:,}",
        f"- Suggested batch size: {budget.suggested_batch_size}",
        f"- Suggested base steps: {budget.suggested_base_steps}",
        f"- Estimated tokens per step: {budget.estimated_tokens_per_step:,}",
        f"- Rough passes over text: {budget.estimated_passes:.2f}",
        f"- Note: {budget.note}",
        "",
        "## Suggested Run Command",
        "",
        f"- Output run: `{report.training_command.out_dir}`",
        f"- Dataset pack: `{report.training_command.dataset_pack}`" if report.training_command.dataset_pack else "- Dataset pack: none",
        f"- Chat SFT input: `{report.training_command.chat_input}`",
        f"- Eval input: `{report.training_command.eval_input}`",
        f"- Note: {report.training_command.note}",
        "",
        "```bash",
        report.training_command.command or f"# {report.training_command.note}",
        "```",
        "",
        "## Chat/Eval Data Preflight",
        "",
        f"- Chat SFT: `{report.chat_data.status}` - {report.chat_data.summary}",
        f"- Chat rows: {report.chat_data.num_examples} usable / {report.chat_data.num_rows} non-empty",
        f"- Eval: `{report.eval_data.status}` - {report.eval_data.summary}",
        f"- Eval rows: {report.eval_data.num_items} usable / {report.eval_data.num_rows} non-empty",
        f"- Eval rules: {report.eval_data.must_include_rules} include, "
        f"{report.eval_data.must_include_any_groups} include-any groups, "
        f"{report.eval_data.must_not_include_rules} forbidden",
        f"- Eval categories: {_format_counts(report.eval_data.categories)}" if report.eval_data.categories else "- Eval categories: none",
        f"- Eval splits: {_format_counts(report.eval_data.splits)}" if report.eval_data.splits else "- Eval splits: none",
        f"- Eval levels: {_format_counts(report.eval_data.levels)}" if getattr(report.eval_data, "levels", {}) else "- Eval levels: none",
        "",
        "## Warnings",
        "",
    ])
    if report.warnings:
        lines.extend(f"- {warning}" for warning in report.warnings)
    else:
        lines.append("- none")
    lines.extend([
        "",
        "## Documents",
        "",
        "These spans let training hold out complete documents instead of random token windows.",
        "",
        "| ID | Source | Label | Chars | Lines | Span | Score |",
        "| ---: | --- | --- | ---: | ---: | --- | ---: |",
    ])
    if report.documents:
        for document in report.documents[:CORPUS_REPORT_MAX_DOCUMENT_ROWS]:
            lines.append(
                f"| {document.document_id} | `{document.path}` | `{document.label or ''}` | "
                f"{document.num_characters} | {document.num_lines} | "
                f"{document.char_start}:{document.char_end} | {document.quality_score} |"
            )
        omitted_documents = len(report.documents) - CORPUS_REPORT_MAX_DOCUMENT_ROWS
        if omitted_documents > 0:
            lines.append(
                f"|  | {omitted_documents:,} more document(s) omitted; see `corpus_manifest.json` for the full list. |  |  |  |  |  |"
            )
    else:
        lines.append("|  | none |  |  |  |  |  |")
    lines.extend([
        "",
        "## Files",
        "",
        "| File | Label | Ext | Included | Score | Chars | Lines | Reason | Flags |",
        "| --- | --- | --- | ---: | ---: | ---: | ---: | --- | --- |",
    ])
    for record in report.files[:CORPUS_REPORT_MAX_FILE_ROWS]:
        flags = ", ".join(record.quality_flags)
        lines.append(
            f"| `{record.path}` | `{record.label or ''}` | `{record.extension}` | {str(record.included).lower()} | "
            f"{record.quality_score} | {record.num_characters} | {record.num_lines} | `{record.reason}` | `{flags}` |"
        )
    omitted_files = len(report.files) - CORPUS_REPORT_MAX_FILE_ROWS
    if omitted_files > 0:
        lines.append(
            f"| {omitted_files:,} more file(s) omitted; see `corpus_manifest.json` for the full list. |  |  |  |  |  |  |  |  |"
        )
    lines.append("")
    return "\n".join(lines)


def _format_counts(counts: dict[str, int]) -> str:
    return ", ".join(f"{name}={count}" for name, count in sorted(counts.items()))


def _document_records(
    sources: tuple[CorpusFileRecord, ...],
    documents: tuple[str, ...],
) -> tuple[CorpusDocumentRecord, ...]:
    records: list[CorpusDocumentRecord] = []
    offset = 0
    for index, (source, text) in enumerate(zip(sources, documents)):
        char_start = offset
        char_end = char_start + len(text)
        records.append(CorpusDocumentRecord(
            document_id=index,
            path=source.path,
            label=source.label,
            char_start=char_start,
            char_end=char_end,
            num_characters=len(text),
            num_lines=len(text.splitlines()),
            quality_score=source.quality_score,
            quality_flags=source.quality_flags,
        ))
        offset = char_end + 2
    return tuple(records)


def _write_corpus_documents(output_path: Path, documents: tuple[str, ...]) -> None:
    with output_path.open("w", encoding="utf-8") as handle:
        for index, document in enumerate(documents):
            if index:
                handle.write("\n\n")
            handle.write(document)
        if documents:
            handle.write("\n")


def _preview_documents(documents: tuple[str, ...], preview_chars: int) -> str:
    if preview_chars <= 0:
        return ""
    remaining = preview_chars
    parts: list[str] = []
    for index, document in enumerate(documents):
        if index:
            separator = "\n\n"
            if remaining <= len(separator):
                parts.append(separator[:remaining])
                break
            parts.append(separator)
            remaining -= len(separator)
        if remaining <= len(document):
            parts.append(document[:remaining])
            break
        parts.append(document)
        remaining -= len(document)
    return "".join(parts)


def _source_files(path: Path) -> list[Path]:
    if path.is_file():
        return [path]
    if not path.exists():
        raise FileNotFoundError(path)
    return sorted(item for item in path.rglob("*") if item.is_file())


def _collect_corpus_sources(
    input_path: str | Path | None,
    recipe_path: str | Path | None = None,
    chat_input: str | Path | None = None,
    eval_input: str | Path | None = None,
    dataset_pack: str | Path | None = None,
    min_quality_score: int = 0,
) -> _CollectedCorpus:
    min_quality_score = _clamp_quality_score(min_quality_score)
    pack = load_dataset_pack(dataset_pack) if dataset_pack else None
    input_path, recipe_path, chat_input, eval_input = _apply_dataset_pack(
        pack,
        input_path,
        recipe_path,
        chat_input,
        eval_input,
    )
    recipe = Path(recipe_path) if recipe_path else None
    input_root = Path(input_path) if input_path else None
    if recipe is None and input_root is None:
        raise ValueError("Either input_path, recipe_path, or dataset_pack is required.")

    input_display = str(input_root) if input_root is not None else str(recipe)
    candidates = _recipe_source_candidates(recipe) if recipe else _path_source_candidates(input_root)
    documents: list[str] = []
    document_sources: list[CorpusFileRecord] = []
    records: list[CorpusFileRecord] = []
    for candidate in candidates:
        text, record = _read_source_candidate(candidate, min_quality_score=min_quality_score)
        records.append(record)
        if text is not None:
            documents.append(text)
            document_sources.append(record)

    stats = inspect_documents(documents, num_files=len(candidates))
    readiness = assess_corpus_readiness(stats, records)
    budget = estimate_training_budget(stats)
    training_command = suggest_training_command(
        input_display,
        str(recipe) if recipe else None,
        budget,
        chat_input=chat_input,
        eval_input=eval_input,
        dataset_pack=str(dataset_pack) if dataset_pack else None,
        min_quality_score=min_quality_score,
    )
    chat_data = inspect_chat_sft_data(training_command.chat_input)
    eval_data = inspect_chat_eval_data(training_command.eval_input)
    warnings = _corpus_warnings(stats, records)
    return _CollectedCorpus(
        input_path=input_display,
        recipe_path=str(recipe) if recipe else None,
        dataset_pack=str(dataset_pack) if dataset_pack else None,
        documents=tuple(documents),
        document_sources=tuple(document_sources),
        files=tuple(records),
        stats=stats,
        readiness=readiness,
        budget=budget,
        training_command=training_command,
        chat_data=chat_data,
        eval_data=eval_data,
        warnings=tuple(warnings),
        min_quality_score=min_quality_score,
    )


def _apply_dataset_pack(
    pack: DatasetPack | None,
    input_path: str | Path | None,
    recipe_path: str | Path | None,
    chat_input: str | Path | None,
    eval_input: str | Path | None,
) -> tuple[str | Path | None, str | Path | None, str | Path | None, str | Path | None]:
    if pack is None:
        return input_path, recipe_path, chat_input, eval_input
    if input_path or recipe_path:
        raise ValueError("Dataset pack cannot be combined with input_path or recipe_path.")
    return pack.corpus_input, pack.corpus_recipe, pack.chat_input, pack.eval_input


def _read_source_candidate(
    candidate: _SourceCandidate,
    min_quality_score: int = 0,
) -> tuple[str | None, CorpusFileRecord]:
    file_path = candidate.path
    extension = file_path.suffix.lower()
    if candidate.skipped_reason:
        return None, CorpusFileRecord(
            path=str(file_path),
            extension=extension or "(none)",
            num_characters=0,
            num_lines=0,
            included=False,
            reason=candidate.skipped_reason,
            quality_score=0,
            quality_flags=(candidate.skipped_reason,),
            label=candidate.label,
        )

    if extension not in SUPPORTED_EXTENSIONS:
        return None, CorpusFileRecord(
            path=str(file_path),
            extension=extension or "(none)",
            num_characters=0,
            num_lines=0,
            included=False,
            reason="unsupported_extension",
            quality_score=0,
            quality_flags=("unsupported_extension",),
            label=candidate.label,
        )

    try:
        text, included_reason = _extract_source_text(file_path)
    except DocumentExtractionError as error:
        return None, CorpusFileRecord(
            path=str(file_path),
            extension=extension,
            num_characters=0,
            num_lines=0,
            included=False,
            reason=error.reason,
            quality_score=0,
            quality_flags=(error.reason,),
            label=candidate.label,
        )

    text = text.strip()
    if not text:
        empty_reason = "empty_text" if extension in TEXT_EXTENSIONS else "empty_extracted_text"
        return None, CorpusFileRecord(
            path=str(file_path),
            extension=extension,
            num_characters=0,
            num_lines=0,
            included=False,
            reason=empty_reason,
            quality_score=0,
            quality_flags=(empty_reason,),
            label=candidate.label,
        )

    quality_score, quality_flags = _score_source_text(text, included_reason)
    if quality_score < min_quality_score:
        return None, CorpusFileRecord(
            path=str(file_path),
            extension=extension,
            num_characters=len(text),
            num_lines=len(text.splitlines()),
            included=False,
            reason="below_min_score",
            quality_score=quality_score,
            quality_flags=(*quality_flags, "below_min_score"),
            label=candidate.label,
        )

    return text, CorpusFileRecord(
        path=str(file_path),
        extension=extension,
        num_characters=len(text),
        num_lines=len(text.splitlines()),
        included=True,
        reason=included_reason,
        quality_score=quality_score,
        quality_flags=quality_flags,
        label=candidate.label,
    )


def _score_source_text(text: str, included_reason: str) -> tuple[int, tuple[str, ...]]:
    """Score one source with explainable local heuristics."""
    lines = text.splitlines()
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    duplicate_lines = len(non_empty_lines) - len(set(non_empty_lines))
    duplicate_rate = duplicate_lines / max(1, len(non_empty_lines))
    empty_rate = sum(1 for line in lines if not line.strip()) / max(1, len(lines))
    non_ascii_rate = sum(1 for char in text if ord(char) > 127) / max(1, len(text))
    control_chars = sum(1 for char in text if ord(char) < 32 and char not in "\n\r\t")

    score = 100
    flags: list[str] = []
    if len(text) < 200:
        score -= 25
        flags.append("short_document")
    elif len(text) < 1000:
        score -= 10
        flags.append("small_document")
    if len(non_empty_lines) <= 1 and len(text) > 500:
        score -= 10
        flags.append("long_single_line")
    if duplicate_rate > 0.30:
        score -= 25
        flags.append("high_duplicate_lines")
    elif duplicate_rate > 0.10:
        score -= 10
        flags.append("duplicate_lines")
    if empty_rate > 0.50:
        score -= 10
        flags.append("many_empty_lines")
    if non_ascii_rate > 0.10:
        score -= 15
        flags.append("high_non_ascii")
    elif non_ascii_rate > 0.05:
        score -= 8
        flags.append("non_ascii")
    if control_chars or "\ufffd" in text:
        score -= 25
        flags.append("extraction_noise")
    if included_reason in {"included_pdf", "included_docx"} and len(text) < 200:
        flags.append("short_extraction")

    return max(0, min(100, score)), tuple(flags)


def _clamp_quality_score(value: int) -> int:
    score = int(value)
    if score < 0 or score > 100:
        raise ValueError("min_quality_score must be between 0 and 100.")
    return score


def _path_source_candidates(path: Path | None) -> list[_SourceCandidate]:
    if path is None:
        return []
    return [_SourceCandidate(item) for item in _source_files(path)]


def _recipe_source_candidates(recipe_path: Path) -> list[_SourceCandidate]:
    recipe = _load_recipe(recipe_path)
    recipe_dir = recipe_path.parent
    global_excludes = _string_list(recipe.get("exclude", []), "exclude")
    sources = recipe.get("sources")
    if not isinstance(sources, list) or not sources:
        raise ValueError("Corpus recipe must define a non-empty 'sources' list.")

    candidates: list[_SourceCandidate] = []
    for index, source in enumerate(sources):
        path_value, label, include, local_excludes, reason = _parse_recipe_source(source, index)
        source_path = Path(path_value)
        if not source_path.is_absolute():
            source_path = recipe_dir / source_path

        if not include:
            candidates.extend(_excluded_candidates(source_path, label, reason))
            continue

        if not source_path.exists():
            candidates.append(_SourceCandidate(source_path, label=label, skipped_reason="missing_source"))
            continue

        excludes = tuple(global_excludes + local_excludes)
        for file_path in _source_files(source_path):
            if _matches_any_pattern(file_path, recipe_dir, excludes):
                candidates.append(_SourceCandidate(file_path, label=label, skipped_reason="recipe_excluded"))
            else:
                candidates.append(_SourceCandidate(file_path, label=label))
    return candidates


def _load_recipe(recipe_path: Path) -> dict:
    try:
        recipe = json.loads(recipe_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        raise ValueError(f"Invalid corpus recipe JSON: {error}") from error
    if not isinstance(recipe, dict):
        raise ValueError("Corpus recipe must be a JSON object.")
    return recipe


def _parse_recipe_source(source: object, index: int) -> tuple[str, str | None, bool, list[str], str]:
    if isinstance(source, str):
        return source, None, True, [], "recipe_excluded"
    if not isinstance(source, dict):
        raise ValueError(f"Corpus recipe source {index} must be a string or object.")

    path_value = source.get("path")
    if not isinstance(path_value, str) or not path_value:
        raise ValueError(f"Corpus recipe source {index} needs a non-empty 'path'.")

    label = source.get("label")
    if label is not None and not isinstance(label, str):
        raise ValueError(f"Corpus recipe source {index} label must be a string.")

    include = source.get("include", True)
    if not isinstance(include, bool):
        raise ValueError(f"Corpus recipe source {index} include must be true or false.")

    excludes = _string_list(source.get("exclude", []), f"sources[{index}].exclude")
    reason = source.get("reason", "recipe_excluded")
    if not isinstance(reason, str) or not reason:
        raise ValueError(f"Corpus recipe source {index} reason must be a non-empty string.")
    return path_value, label, include, excludes, reason


def _string_list(value: object, field_name: str) -> list[str]:
    if value is None:
        return []
    if not isinstance(value, list) or not all(isinstance(item, str) for item in value):
        raise ValueError(f"Corpus recipe field '{field_name}' must be a list of strings.")
    return value


def _excluded_candidates(path: Path, label: str | None, reason: str) -> list[_SourceCandidate]:
    if path.exists() and path.is_dir():
        return [_SourceCandidate(file_path, label=label, skipped_reason=reason) for file_path in _source_files(path)]
    return [_SourceCandidate(path, label=label, skipped_reason=reason)]


def _matches_any_pattern(file_path: Path, recipe_dir: Path, patterns: tuple[str, ...]) -> bool:
    if not patterns:
        return False
    relative_path = _relative_posix(file_path, recipe_dir)
    name = file_path.name
    return any(fnmatch(relative_path, pattern) or fnmatch(name, pattern) for pattern in patterns)


def _relative_posix(file_path: Path, root: Path) -> str:
    try:
        return file_path.relative_to(root).as_posix()
    except ValueError:
        return file_path.as_posix()


def _extract_source_text(file_path: Path) -> tuple[str, str]:
    extension = file_path.suffix.lower()
    if extension in TEXT_EXTENSIONS:
        return file_path.read_text(encoding="utf-8", errors="replace"), "included"
    if extension == ".pdf":
        return _extract_pdf_text(file_path), "included_pdf"
    if extension == ".docx":
        return _extract_docx_text(file_path), "included_docx"
    raise DocumentExtractionError("unsupported_extension")


def _extract_pdf_text(file_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as import_error:
        raise DocumentExtractionError("missing_pdf_dependency") from import_error

    try:
        reader = PdfReader(str(file_path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as error:
        raise DocumentExtractionError(f"extract_failed:{error.__class__.__name__}") from error
    return "\n\n".join(page for page in pages if page)


def _extract_docx_text(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as import_error:
        raise DocumentExtractionError("missing_docx_dependency") from import_error

    try:
        document = Document(str(file_path))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as error:
        raise DocumentExtractionError(f"extract_failed:{error.__class__.__name__}") from error
    return "\n".join(parts)


def _corpus_warnings(stats: CorpusStats, records: list[CorpusFileRecord]) -> list[str]:
    warnings: list[str] = []
    skipped = [record for record in records if not record.included]
    filtered = [record for record in skipped if record.reason == "below_min_score"]
    low_quality = [
        record for record in records
        if record.included and record.quality_score < 70
    ]
    if stats.num_documents == 0:
        warnings.append("No usable text documents were included.")
    if stats.num_characters < 1000:
        warnings.append("Corpus is very small; expect memorization and weak generalization.")
    if stats.num_documents == 1:
        warnings.append("Only one usable document was included; add more sources for broader variation.")
    if stats.duplicate_line_rate > 0.15:
        warnings.append("Duplicate line rate is high; repeated text can encourage memorization.")
    if stats.duplicate_document_rate > 0.05:
        warnings.append("Duplicate document rate is high; deduplicate sources before a serious run.")
    if stats.near_duplicate_document_rate > 0.05:
        warnings.append("Near-duplicate document rate is high; add dedup or sampling before scaling.")
    if stats.empty_line_rate > 0.35:
        warnings.append("Empty line rate is high; context windows may be wasted.")
    if stats.non_ascii_rate > 0.05:
        warnings.append("Non-ASCII rate is high; confirm this is expected for the corpus.")
    if low_quality:
        warnings.append(f"{len(low_quality)} included source file(s) scored below 70.")
    if filtered:
        warnings.append(f"{len(filtered)} source file(s) were filtered by the minimum quality score.")
    if skipped:
        warnings.append(f"{len(skipped)} source file(s) were skipped.")
    return warnings
